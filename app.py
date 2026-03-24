import streamlit as st
import io
import os
import tempfile
from datetime import datetime

from convert import convert_md_to_docx
from file_converter import convert_file_to_md


def download_template_from_drive(file_id: str) -> str:
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    from google.oauth2 import service_account
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = service_account.Credentials.from_service_account_info(
        creds_dict, scopes=["https://www.googleapis.com/auth/drive.readonly"])
    service = build("drive", "v3", credentials=creds)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    request = service.files().get_media(fileId=file_id)
    downloader = MediaIoBaseDownload(tmp, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    tmp.close()
    return tmp.name


def get_template(use_drive, drive_file_id, local_path):
    if use_drive:
        try:
            with st.spinner("Загружаю шаблон с Google Drive..."):
                path = download_template_from_drive(drive_file_id)
            return path
        except Exception as e:
            st.warning(f"⚠️ Не удалось загрузить шаблон с Drive: {e}")
    if local_path and os.path.exists(local_path):
        return local_path
    return None


DOC_TYPES = {
    "📄 Письмо / Сопроводительное письмо": {
        "drive_id":    "1FdPo8Ddo317ZYoPzraCTy5R4E72Ieqba",
        "local_path":  r"C:\Users\tonik\Desktop\docx_converter\template.docx",
        "output_name": "letter",
        "hint": "Структура: заголовок `# Название`, блок `**Кому:**`, разделы `## ...`, подпись `С уважением,`"
    },
    "📋 Пояснительная записка": {
        "drive_id":    "1FdPo8Ddo317ZYoPzraCTy5R4E72Ieqba",
        "local_path":  r"C:\Users\tonik\Desktop\docx_converter\template.docx",
        "output_name": "pz",
        "hint": "Структура: заголовок `# Название`, разделы `## 1. ...`, подразделы `### 1.1. ...`, callout `!! формула !!`"
    },
}


def insert_images_into_docx(docx_path, images):
    if not images:
        with open(docx_path, 'rb') as f:
            return f.read()
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:color'), 'CCCCCC')
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run('Изображения из исходного документа')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    for fname, img_bytes in images:
        try:
            tmp_img = tempfile.NamedTemporaryFile(
                delete=False, suffix='.' + fname.split('.')[-1])
            tmp_img.write(img_bytes)
            tmp_img.close()
            p2 = doc.add_paragraph()
            p2.add_run().add_picture(tmp_img.name, width=Cm(14))
            os.unlink(tmp_img.name)
        except Exception:
            pass
    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_out.close()
    doc.save(tmp_out.name)
    with open(tmp_out.name, 'rb') as f:
        result = f.read()
    os.unlink(tmp_out.name)
    return result


st.set_page_config(
    page_title="Тензосила — Конструктор документов",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    .block-container { padding-top: 2rem; padding-bottom: 2rem; }
    .stTextArea textarea { font-family: monospace; font-size: 13px; }
</style>
""", unsafe_allow_html=True)

col_logo, col_title = st.columns([1, 5])
with col_logo:
    st.markdown("## ⚖️")
with col_title:
    st.markdown("## Конструктор документов")
    st.caption("ООО «ТПК «Тензосила» · Фирменное оформление по брендбуку")

st.divider()

col_left, col_right = st.columns([3, 2], gap="large")

with col_left:
    st.markdown("#### 1. Тип документа")
    doc_type = st.selectbox("Тип", options=list(DOC_TYPES.keys()),
                            label_visibility="collapsed")
    config = DOC_TYPES[doc_type]

    with st.expander("📌 Как оформить текст для этого типа", expanded=False):
        st.code(config["hint"], language=None)

    st.markdown("#### 2. Текст документа")

    tab_paste, tab_md, tab_file = st.tabs([
        "✏️ Вставить текст",
        "📂 Загрузить .md",
        "📄 Загрузить файл (docx / pdf / txt)"
    ])

    md_text = ""
    source_images = []

    with tab_paste:
        md_input = st.text_area(
            "Markdown текст", height=400,
            placeholder="# Заголовок\n\nТекст...",
            label_visibility="collapsed")
        if md_input:
            md_text = md_input

    with tab_md:
        upl_md = st.file_uploader("MD файл", type=["md", "txt"],
                                  label_visibility="collapsed", key="upl_md")
        if upl_md:
            md_text = upl_md.read().decode("utf-8")
            st.success(f"Загружен: **{upl_md.name}** · {len(md_text)} символов")
            with st.expander("👁 Превью", expanded=False):
                st.text(md_text[:1500] + ("..." if len(md_text) > 1500 else ""))

    with tab_file:
        st.caption(
            "Загрузи DOCX, PDF или TXT — скрипт определит заголовки "
            "по размеру и жирности шрифта и переоформит в фирменный стиль."
        )
        upl_file = st.file_uploader(
            "Файл", type=["docx", "pdf", "txt"],
            label_visibility="collapsed", key="upl_file")

        if upl_file:
            file_bytes = upl_file.read()
            with st.spinner(f"Извлекаю текст из {upl_file.name}..."):
                try:
                    md_text, source_images = convert_file_to_md(
                        file_bytes, upl_file.name)

                    img_info = (f" · {len(source_images)} изображений"
                                if source_images else "")
                    st.success(
                        f"✅ Обработан: **{upl_file.name}** · "
                        f"{len(md_text)} символов{img_info}")

                    with st.expander("👁 Распознанная структура", expanded=False):
                        headers = [l for l in md_text.split('\n')
                                   if l.startswith('#')]
                        st.text('\n'.join(headers[:30]) if headers
                                else md_text[:800])

                    edited = st.text_area(
                        "✏️ Отредактируй при необходимости:",
                        value=md_text, height=300, key="edit_md")
                    if edited:
                        md_text = edited

                except Exception as e:
                    st.error(f"❌ Ошибка: {e}")
                    md_text = ""

with col_right:
    st.markdown("#### 3. Сформировать документ")

    default_name = f"{config['output_name']}_{datetime.now().strftime('%d%m%Y')}"
    output_name  = st.text_input("Имя файла (без расширения)",
                                  value=default_name)
    st.markdown("---")

    btn = st.button("⚙️ Сформировать документ", type="primary",
                    use_container_width=True,
                    disabled=(not md_text.strip()))

    if not md_text.strip():
        st.caption("Вставь текст или загрузи файл")

    if btn and md_text.strip():
        with st.spinner("Формирую документ..."):
            use_drive = "gcp_service_account" in st.secrets
            template_path = get_template(
                use_drive=use_drive,
                drive_file_id=config["drive_id"],
                local_path=config["local_path"])

            if template_path is None:
                st.error("❌ Шаблон не найден.")
                st.stop()

            try:
                tmp_out = tempfile.NamedTemporaryFile(
                    delete=False, suffix=".docx")
                tmp_out.close()

                convert_md_to_docx(md_text=md_text,
                                   output_filename=tmp_out.name,
                                   template_path=template_path)

                docx_bytes = insert_images_into_docx(
                    tmp_out.name, source_images)

                os.unlink(tmp_out.name)
                if use_drive and template_path != config["local_path"]:
                    try:
                        os.unlink(template_path)
                    except Exception:
                        pass

                st.success("✅ Документ готов!")
                st.download_button(
                    "⬇️ Скачать .docx", data=docx_bytes,
                    file_name=f"{output_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document",
                    use_container_width=True)
                st.caption(f"Размер: {len(docx_bytes)/1024:.1f} КБ")

            except Exception as e:
                st.error(f"❌ Ошибка:\n\n```\n{e}\n```")

    st.markdown("---")
    st.markdown("#### 📖 Шпаргалка")
    st.markdown("""
| Что | Разметка |
|-----|----------|
| Заголовок | `# Текст` |
| Раздел | `## Текст` |
| Подраздел | `### Текст` |
| **Жирный** | `**текст**` |
| Таблица | `\\| А \\| Б \\|` |
| Врезка | `!! текст !!` |
| Блок «Кому» | `**Кому:** ...` |
| Подпись | `С уважением,` |
""")

st.divider()
st.caption("ООО «ТПК «Тензосила» · Внутренний инструмент")
