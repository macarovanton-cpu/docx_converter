import os
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# НАСТРОЙКИ — МЕНЯЙ ТОЛЬКО ЭТИ ТРИ СТРОКИ
# =============================================================================
INPUT_FILE    = r"C:\Users\tonik\Desktop\docx_converter\Cheremhovo.md"
OUTPUT_FILE   = r"C:\Users\tonik\Desktop\docx_converter\Cheremhovo.docx"
TEMPLATE_FILE = r"C:\Users\tonik\Desktop\docx_converter\template.docx"
# =============================================================================

# === ЦВЕТОВАЯ ПАЛИТРА ===
BRAND_BLUE    = "015198"
BRAND_RED     = "D04514"
BRAND_WHITE   = "FFFFFF"
TEXT_DARK     = "1A1A1A"
TEXT_MUTED    = "888888"
BG_LIGHT_BLUE = "EBF3FB"
BG_TABLE_ROW  = "F0F5FA"
BG_TABLE_LAST = "EBF3FB"  # последняя колонка в сравнительных таблицах
BORDER_LIGHT  = "CCCCCC"
COLOR_YES     = "1E7A34"
COLOR_NO      = "C0392B"

# Поля шаблона: left=2cm, right=1.5cm
CONTENT_WIDTH_CM = 17.5


# =============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ — те же что в convert.py
# =============================================================================

def clear_body(doc):
    body = doc.element.body
    to_remove = [c for c in body
                 if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) != 'sectPr']
    for el in to_remove:
        body.remove(el)
    body.insert(0, OxmlElement('w:p'))


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_margins_and_borders(cell, hex_color, sz):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, w in [('top','100'),('bottom','100'),('left','160'),('right','160')]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), w)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), str(sz))
        bdr.set(qn('w:color'), hex_color)
        tcBorders.append(bdr)
    tcPr.append(tcBorders)


def add_paragraph_border(paragraph, side, color, size, space="0"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), str(space))
    border.set(qn('w:color'), color)
    pBdr.append(border)


def add_paragraph_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_table_width_dxa(table, width_cm):
    dxa = int(width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_row_height(row, height_dxa):
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        row._tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_dxa))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))


def set_keep_together(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))


def set_run_font(run, name, size_pt, color_hex, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def add_page_number_field(run):
    for ftype in ['begin', None, 'separate', 'end']:
        if ftype is None:
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = 'PAGE'
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), ftype)
        run._r.append(el)


def parse_inline_markdown(paragraph, text, font_name='PT Sans', font_size=12,
                          font_color=TEXT_DARK, is_italic_base=False):
    pattern = re.compile(r'(\*\*[^*\n]+?\*\*|\*(?!\*)[^*\n]+?\*(?!\*))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        is_bold   = False
        is_italic = is_italic_base
        clean_text = part
        if part.startswith('**') and part.endswith('**') and len(part) >= 5:
            clean_text = part[2:-2]
            is_bold = True
        elif (part.startswith('*') and part.endswith('*')
              and len(part) >= 3 and not part.startswith('**')):
            clean_text = part[1:-1]
            is_italic = True
        set_run_font(run, font_name, font_size, font_color,
                     bold=is_bold, italic=is_italic)
        run.text = clean_text


def add_table_cell_content(p, text, font_size=10):
    """Добавляет ✓/✗ для Да/Нет/Отсутствует в ячейках."""
    stripped = text.strip()
    if re.match(r'^Да\b', stripped, re.IGNORECASE):
        r = p.add_run('✓ ')
        set_run_font(r, 'PT Sans', font_size, COLOR_YES, bold=True)
        rest = re.sub(r'^Да\b\s*', '', stripped, flags=re.IGNORECASE)
        if rest:
            parse_inline_markdown(p, rest, 'PT Sans', font_size, TEXT_DARK)
        else:
            p.add_run('Да').font.color.rgb = RGBColor.from_string(TEXT_DARK)
    elif re.match(r'^Нет\b', stripped, re.IGNORECASE):
        r = p.add_run('✗ ')
        set_run_font(r, 'PT Sans', font_size, COLOR_NO, bold=True)
        rest = re.sub(r'^Нет\b\s*', '', stripped, flags=re.IGNORECASE)
        if rest:
            parse_inline_markdown(p, rest, 'PT Sans', font_size, TEXT_DARK)
        else:
            p.add_run('Нет').font.color.rgb = RGBColor.from_string(TEXT_DARK)
    elif re.match(r'^Отсутствует\b', stripped, re.IGNORECASE):
        r = p.add_run('✗ ')
        set_run_font(r, 'PT Sans', font_size, COLOR_NO, bold=True)
        rest = re.sub(r'^Отсутствует\b\s*', '', stripped, flags=re.IGNORECASE)
        parse_inline_markdown(p, ('Отсутствует ' + rest).strip(), 'PT Sans', font_size, TEXT_DARK)
    else:
        parse_inline_markdown(p, stripped, 'PT Sans', font_size, TEXT_DARK)


def is_callout_block(text):
    return text.startswith('!!') and text.endswith('!!')


# =============================================================================
# ПАРСИНГ YAML FRONTMATTER
# =============================================================================

def parse_frontmatter(md_text):
    """
    Читает YAML-шапку между --- разделителями в начале файла.
    Возвращает (metadata_dict, text_without_frontmatter).
    """
    meta = {}
    lines = md_text.split('\n')

    if lines[0].strip() != '---':
        return meta, md_text

    end_idx = None
    for i, line in enumerate(lines[1:], 1):
        if line.strip() == '---':
            end_idx = i
            break

    if end_idx is None:
        return meta, md_text

    for line in lines[1:end_idx]:
        if ':' in line:
            key, _, value = line.partition(':')
            meta[key.strip()] = value.strip()

    remaining = '\n'.join(lines[end_idx + 1:]).strip()
    return meta, remaining


# =============================================================================
# ТИТУЛЬНЫЙ БЛОК ПЗ
# =============================================================================

def add_title_block(doc, meta, content_width_cm):
    """
    Рендерит шапку документа из метаданных YAML frontmatter.
    Структура: название документа, клиент, оборудование, номер, дата, автор.
    """
    title    = meta.get('title', 'Пояснительная записка')
    client   = meta.get('client', '')
    equip    = meta.get('equipment', '')
    doc_num  = meta.get('doc_number', '')
    date     = meta.get('date', '')
    author   = meta.get('author', '')

    # Заголовок документа
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after  = Pt(6)
    r = p_title.add_run(title)
    set_run_font(r, 'PT Sans Narrow', 18, BRAND_BLUE, bold=True)

    # Разделитель под заголовком
    dec = doc.add_paragraph()
    dec.paragraph_format.space_before = Pt(0)
    dec.paragraph_format.space_after  = Pt(10)
    add_paragraph_border(dec, 'bottom', BRAND_RED, 12)

    # Мета-строки: клиент, оборудование, номер, дата — в одном блоке с фоном
    meta_lines = []
    if client:  meta_lines.append(('Заказчик:', client))
    if equip:   meta_lines.append(('Оборудование:', equip))
    if doc_num: meta_lines.append(('№ документа:', doc_num))
    if date:    meta_lines.append(('Дата:', date))
    if author:  meta_lines.append(('Составил:', author))

    if meta_lines:
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
        p_meta.paragraph_format.space_before = Pt(2)
        p_meta.paragraph_format.space_after  = Pt(14)
        p_meta.paragraph_format.left_indent  = Cm(0.4)
        add_paragraph_shading(p_meta, BG_LIGHT_BLUE)

        for i, (label, value) in enumerate(meta_lines):
            if i > 0:
                p_meta.add_run().add_break()
            r_label = p_meta.add_run(label + ' ')
            set_run_font(r_label, 'PT Sans', 11, TEXT_DARK, bold=True)
            r_value = p_meta.add_run(value)
            set_run_font(r_value, 'PT Sans', 11, TEXT_DARK)


# =============================================================================
# БЛОК CALLOUT (врезка с формулой или ключевым выводом)
# =============================================================================

def add_callout_box(doc, text, content_width_cm):
    clean = text.strip('!').strip()
    table = doc.add_table(rows=1, cols=1)
    table.autofit = table.allow_autofit = False
    set_table_width_dxa(table, content_width_cm)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F2F6FA')

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:color'), 'C5D8EC')
        tcBorders.append(bdr)
    tcPr.append(tcBorders)
    tcMar = OxmlElement('w:tcMar')
    for side, w in [('top','140'),('bottom','140'),('left','220'),('right','220')]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), w)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    parse_inline_markdown(p, clean, 'PT Sans', 12, BRAND_BLUE)
    for r in p.runs:
        r.bold = True


# =============================================================================
# ОСНОВНАЯ ЛОГИКА
# =============================================================================

def convert_pz(md_text, output_filename, template_path=None):

    # Читаем метаданные из YAML шапки
    meta, body_text = parse_frontmatter(md_text)

    # Открываем шаблон или создаём чистый документ
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        clear_body(doc)
        content_width_cm = CONTENT_WIDTH_CM
        print(f"  Шаблон: {template_path}")
    else:
        doc = Document()
        section = doc.sections[0]
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        content_width_cm = 21.0 - 2.54 * 2
        print("  ⚠️  Шаблон не найден — создаётся без хедера")

        footer   = doc.sections[0].footer
        footer_p = footer.paragraphs[0]
        footer_p.paragraph_format.tab_stops.add_tab_stop(
            Cm(content_width_cm), WD_TAB_ALIGNMENT.RIGHT)
        rl = footer_p.add_run("ООО «ТПК «Тензосила»")
        set_run_font(rl, 'PT Sans', 9, TEXT_MUTED)
        footer_p.add_run("\t")
        rr = footer_p.add_run()
        set_run_font(rr, 'PT Sans', 9, TEXT_MUTED)
        add_page_number_field(rr)

    # Базовый стиль Normal
    sn = doc.styles['Normal']
    sn.font.name      = 'PT Sans'
    sn.font.size      = Pt(12)
    sn.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    sn.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    sn.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sn.paragraph_format.line_spacing      = 1.3
    sn.paragraph_format.space_after       = Pt(8)

    # Рендерим титульный блок из метаданных
    if meta:
        add_title_block(doc, meta, content_width_cm)

    # Парсим тело документа
    blocks = body_text.split('\n\n')
    after_heading        = False
    last_paragraph       = None

    for block in blocks:
        block = block.strip()
        if not block or block == '---':
            continue

        is_list_item = (block.startswith('- ')
                        or block.startswith('* ')
                        or bool(re.match(r'^\d+\.', block)))

        # ── H1 — в ПЗ обычно не используется как раздел, это подзаголовок ──
        if block.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
            parse_inline_markdown(p, block[2:], 'PT Sans Narrow', 16, BRAND_BLUE)
            for r in p.runs: r.bold = True
            after_heading = True

        # ── H2 — основные разделы ПЗ (1. / 2. / 3.) ───────────────────────
        elif block.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
            # В ПЗ: чёрный bold — строгий технический стиль
            parse_inline_markdown(p, block[3:], 'PT Sans Narrow', 13, TEXT_DARK)
            for r in p.runs: r.bold = True
            set_keep_with_next(p)
            after_heading = True

        # ── H3 — подразделы (1.1. / 2.1.) ──────────────────────────────────
        elif block.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
            # Подраздел: чёрный 11.5pt, чуть меньше H2
            parse_inline_markdown(p, block[4:], 'PT Sans', 11, TEXT_DARK)
            for r in p.runs: r.bold = True
            set_keep_with_next(p)
            after_heading = True

        # ── Callout-врезка !! формула или вывод !! ──────────────────────────
        elif is_callout_block(block):
            add_callout_box(doc, block, content_width_cm)
            after_heading = False

        # ── Списки ──────────────────────────────────────────────────────────
        elif is_list_item:
            for line in block.split('\n'):
                line = line.strip()
                if not line: continue
                is_num = bool(re.match(r'^\d+\.', line))
                p = doc.add_paragraph(style='List Number' if is_num else 'List Bullet')
                p.paragraph_format.left_indent       = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.75)
                p.paragraph_format.space_after       = Pt(4)
                parse_inline_markdown(p, re.sub(r'^(\- |\* |\d+\. )', '', line))
                last_paragraph = p
            after_heading = False

        # ── Таблицы ─────────────────────────────────────────────────────────
        elif block.startswith('|') and '\n|' in block:
            lines = [l.strip() for l in block.split('\n')
                     if l.strip() and not re.match(r'^\|[-| ]+\|$', l.strip())]
            if not lines: continue

            headers  = [c.strip() for c in lines[0].strip('|').split('|')]
            n_cols   = len(headers)
            table    = doc.add_table(rows=1, cols=n_cols)
            table.autofit = table.allow_autofit = False
            set_table_width_dxa(table, content_width_cm)

            # Равномерное распределение для 2 колонок,
            # 30/35/35 для 3 колонок (первая — «Параметр» обычно шире)
            if n_cols == 2:
                ratios = [0.45, 0.55]
            elif n_cols == 3:
                ratios = [0.35, 0.30, 0.35]
            elif n_cols == 4:
                ratios = [0.30, 0.25, 0.25, 0.20]
            else:
                ratios = [1/n_cols] * n_cols

            total_dxa  = int(content_width_cm * 567)
            widths_dxa = [int(content_width_cm * 567 * r) for r in ratios]
            widths_dxa[-1] = total_dxa - sum(widths_dxa[:-1])
            for i, col in enumerate(table.columns):
                col.width = widths_dxa[i]

            # Заголовочная строка
            set_row_height(table.rows[0], 520)
            for i, h in enumerate(headers):
                if i < len(table.rows[0].cells):
                    cell = table.rows[0].cells[i]
                    set_cell_shading(cell, BRAND_BLUE)
                    set_cell_margins_and_borders(cell, BORDER_LIGHT, 4)
                    p = cell.paragraphs[0]
                    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    parse_inline_markdown(p, h, 'PT Sans Narrow', 11, BRAND_WHITE)
                    for r in p.runs: r.bold = True

            # Строки данных
            for row_idx, line in enumerate(lines[1:]):
                cols      = [c.strip() for c in line.strip('|').split('|')]
                row_cells = table.add_row().cells
                bg        = BG_TABLE_ROW if (row_idx % 2 == 1) else 'FFFFFF'

                for i, c in enumerate(cols):
                    if i < len(row_cells):
                        cell = row_cells[i]
                        # Последняя колонка 3+ таблиц — лёгкий акцент
                        cell_bg = BG_TABLE_LAST if (n_cols >= 3 and i == n_cols - 1) else bg
                        set_cell_shading(cell, cell_bg)
                        set_cell_margins_and_borders(cell, BORDER_LIGHT, 4)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(0)
                        add_table_cell_content(p, c, font_size=10)

            # Отступ после таблицы
            sp = doc.add_paragraph()
            pPr = sp._p.get_or_add_pPr()
            s = OxmlElement('w:spacing')
            s.set(qn('w:before'), '0'); s.set(qn('w:after'), '120')
            s.set(qn('w:line'), '120');  s.set(qn('w:lineRule'), 'exact')
            pPr.append(s)
            after_heading = False

        # ── Обычные абзацы ──────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            if not after_heading:
                p.paragraph_format.first_line_indent = Cm(0.75)
            for i, line in enumerate(block.split('\n')):
                line = line.strip()
                if not line: continue
                if i > 0: p.add_run().add_break()
                parse_inline_markdown(p, line)
            last_paragraph = p
            after_heading = False

    doc.save(output_filename)
    print(f"✅ Готово! Файл сохранён: {output_filename}")


# =============================================================================
# ЗАПУСК
# =============================================================================
if __name__ == '__main__':
    if not os.path.exists(INPUT_FILE):
        print(f"❌ .md файл не найден: {INPUT_FILE}")
        print("   Проверь путь INPUT_FILE в начале скрипта.")
        exit(1)

    if not os.path.exists(TEMPLATE_FILE):
        print(f"⚠️  Шаблон не найден: {TEMPLATE_FILE}")
        print("   Документ будет создан без фирменного хедера.\n")

    print(f"Конвертирую ПЗ: {os.path.basename(INPUT_FILE)}")
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        md_text = f.read()

    template = TEMPLATE_FILE if os.path.exists(TEMPLATE_FILE) else None
    convert_pz(md_text, OUTPUT_FILE, template_path=template)
