"""
file_converter.py
Извлекает текст и структуру из DOCX / PDF / TXT
и возвращает:
  - md_text: строка в формате Markdown
  - images: список (filename, bytes) для вставки в итоговый документ
"""

import io
import os
import re
import tempfile
from html import unescape as html_unescape
from typing import Optional


# =============================================================================
# MarkItDown layer
# =============================================================================

PDF_IMAGE_ONLY_MESSAGE = (
    "Возможно, выбранные страницы содержат только изображение. "
    "Для извлечения текста потребуется OCR."
)


def get_pdf_page_count(file_path: str) -> int:
    """Возвращает количество страниц в PDF."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    return len(reader.pages)


def analyze_pdf_pages(file_path: str) -> list[dict]:
    """
    Возвращает диагностику текстового слоя PDF по страницам.

    Формат элемента:
    {
        "page_number": 1,
        "has_text_layer": True,
        "extracted_chars": 123,
        "error": None,
    }
    """
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for index in range(len(reader.pages)):
        page_info = {
            "page_number": index + 1,
            "has_text_layer": False,
            "extracted_chars": 0,
            "error": None,
        }
        try:
            page = reader.pages[index]
            text = page.extract_text() or ''
        except Exception as e:
            text = ''
            page_info["error"] = str(e)
        extracted_chars = len(text.strip())
        page_info["extracted_chars"] = extracted_chars
        page_info["has_text_layer"] = extracted_chars > 0
        pages.append(page_info)
    return pages


def parse_page_range(range_text: str) -> list[int] | None:
    """
    Парсит пользовательский диапазон страниц в 0-based индексы.

    Пример: "1-3, 7, 10-12" -> [0, 1, 2, 6, 9, 10, 11].
    Пустая строка означает "все страницы" и возвращает None.
    """
    if range_text is None:
        return None

    text = range_text.strip()
    if not text:
        return None

    pages = []
    seen = set()

    for raw_part in text.split(','):
        part = raw_part.strip()
        if not part:
            raise ValueError(
                f"Некорректный диапазон страниц: {range_text!r}."
            )

        match = re.fullmatch(r'(\d+)(?:\s*-\s*(\d+))?', part)
        if not match:
            raise ValueError(
                "Некорректный диапазон страниц. Используйте формат "
                "'1-3, 7, 10-12'."
            )

        start = int(match.group(1))
        end = int(match.group(2) or start)

        if start < 1 or end < 1:
            raise ValueError("Номера страниц должны начинаться с 1.")
        if start > end:
            raise ValueError(
                f"Некорректный диапазон страниц: {start}-{end}."
            )

        if end - start + 1 > 1000:
            raise ValueError(
                f"Диапазон '{start}-{end}' содержит {end - start + 1} страниц — "
                "максимально допустимо 1 000."
            )
        for page_num in range(start, end + 1):
            page_index = page_num - 1
            if page_index not in seen:
                pages.append(page_index)
                seen.add(page_index)
        if len(pages) > 1000:
            raise ValueError(
                "Суммарное количество выбранных страниц превышает 1 000. "
                "Разбейте запрос на несколько меньших диапазонов."
            )

    return pages


def convert_with_markitdown(file_path: str, page_range: str | None = None) -> str:
    """
    Конвертирует PDF/DOCX/XLSX/PPTX в Markdown через Microsoft MarkItDown.

    Диапазон страниц на первом этапе поддержан только для PDF. Для PDF
    создаётся временный файл с выбранными 1-based страницами пользователя,
    после чего именно он передаётся в MarkItDown.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    supported_exts = {'.pdf', '.docx', '.xlsx', '.pptx'}
    if ext not in supported_exts:
        raise ValueError(
            f"Неподдерживаемый формат для MarkItDown: {ext or 'без расширения'}."
        )

    has_page_range = bool(page_range and page_range.strip())
    if has_page_range and ext != '.pdf':
        raise ValueError(
            "Диапазон страниц на первом этапе поддержан только для PDF. "
            "Для DOCX/XLSX/PPTX конвертируйте файл целиком."
        )

    pages = parse_page_range(page_range) if has_page_range else None

    source_path = file_path
    tmp_pdf_path = None
    selected_pdf_pages = None

    try:
        if ext == '.pdf' and pages is not None:
            selected_pdf_pages = _get_selected_pdf_page_analysis(file_path, pages)
            source_path = _create_pdf_page_subset(file_path, pages)
            tmp_pdf_path = source_path

        from markitdown import MarkItDown

        result = MarkItDown().convert(source_path)
        markdown = getattr(result, 'markdown', '')
        if (
            ext == '.pdf'
            and selected_pdf_pages is not None
            and _looks_like_image_only_result(markdown, selected_pdf_pages)
        ):
            raise ValueError(PDF_IMAGE_ONLY_MESSAGE)
        if not markdown or not markdown.strip():
            raise ValueError(
                "MarkItDown вернул пустой Markdown. Проверьте, что файл "
                "содержит извлекаемый текст."
            )
        return markdown.strip()
    finally:
        if tmp_pdf_path:
            try:
                os.unlink(tmp_pdf_path)
            except OSError:
                pass


def _get_selected_pdf_page_analysis(file_path: str,
                                    page_indexes: list[int]) -> list[dict]:
    """Возвращает анализ выбранных 0-based страниц и проверяет границы PDF."""
    all_pages = analyze_pdf_pages(file_path)
    total_pages = len(all_pages)
    selected = []

    for page_index in page_indexes:
        if page_index < 0 or page_index >= total_pages:
            raise ValueError(
                f"Страница {page_index + 1} вне диапазона PDF: "
                f"в файле всего {total_pages} стр."
            )
        selected.append(all_pages[page_index])

    return selected


def _looks_like_image_only_result(markdown: str,
                                  selected_pages: list[dict]) -> bool:
    """Определяет, что выбранные страницы, вероятно, image-only."""
    markdown_chars = len((markdown or '').strip())
    has_text_pages = [
        page for page in selected_pages
        if page.get("has_text_layer") or page.get("extracted_chars", 0) > 0
    ]
    return not has_text_pages and markdown_chars < 50


def _create_pdf_page_subset(file_path: str, page_indexes: list[int]) -> str:
    """Создаёт временный PDF с выбранными 0-based страницами."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(file_path)
    total_pages = len(reader.pages)

    writer = PdfWriter()
    for page_index in page_indexes:
        if page_index >= total_pages:
            raise ValueError(
                f"Страница {page_index + 1} вне диапазона PDF: "
                f"в файле всего {total_pages} стр."
            )
        writer.add_page(reader.pages[page_index])

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    try:
        writer.write(tmp)
        return tmp.name
    finally:
        tmp.close()


# =============================================================================
# DOCX → MD
# =============================================================================

def docx_to_md(file_bytes: bytes) -> tuple[str, list]:
    """
    Конвертирует DOCX в Markdown с помощью mammoth.
    Возвращает (md_text, images).
    images = список (image_filename, image_bytes) в порядке появления.
    """
    import mammoth

    style_map = "\n".join([
        "p[style-name='Title'] => h1:fresh",
        "p[style-name='Heading 1'] => h1:fresh",
        "p[style-name='Heading 2'] => h2:fresh",
        "p[style-name='Heading 3'] => h3:fresh",
        "p[style-name='Название'] => h1:fresh",
        "p[style-name='Заголовок 1'] => h1:fresh",
        "p[style-name='Заголовок 2'] => h2:fresh",
        "p[style-name='Заголовок 3'] => h3:fresh",
    ])

    # ПРАВКА #23: собираем картинки через mammoth callback — порядок = позиция в тексте
    images = []
    counter = [0]

    @mammoth.images.img_element
    def _collect_image(image):
        n = counter[0]
        counter[0] += 1
        with image.open() as f:
            img_bytes = f.read()
        ext = image.content_type.split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        fname = f'image_{n}.{ext}'
        images.append((fname, img_bytes))
        return {'src': fname}

    try:
        result = mammoth.convert_to_html(
            io.BytesIO(file_bytes),
            style_map=style_map,
            convert_image=_collect_image,
        )
        md_text = _html_to_md(result.value)
    except Exception as e:
        return (
            f"❌ Не удалось извлечь текст из DOCX: {e}\n\n"
            "Попробуйте сохранить документ заново через Файл → Сохранить как.",
            [],
        )

    md_text = _postprocess_md(md_text)
    return md_text, images


def _extract_images_from_docx(doc) -> list[tuple[str, bytes]]:
    """Извлекает изображения из python-docx Document через rels."""
    images = []
    counter = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_part.content_type.split('/')[-1]
                if ext == 'jpeg':
                    ext = 'jpg'
                fname = f"image_{counter}.{ext}"
                images.append((fname, img_bytes))
                counter += 1
            except Exception:
                pass
    return images


def _html_to_md(html: str) -> str:
    """Конвертирует HTML-вывод mammoth в Markdown."""
    h = html

    # ПРАВКА #23: <img src="..."> → ![alt](src) вместо удаления
    h = re.sub(r'<img\s+src="([^"]*)"(?:\s+alt="([^"]*)")?\s*/?>',
               lambda m: f'![{m.group(2) or ""}]({m.group(1)})', h)

    # Inline-форматирование через плейсхолдеры из непечатаемых символов.
    # Это нужно, чтобы соседние <strong> и <em> не давали каскад звёздочек
    # вида ***...***...*** на стыках, который ломает парсер convert.py.
    BI_O, BI_C = '\x01', '\x02'  # bold-italic
    B_O,  B_C  = '\x03', '\x04'  # bold
    I_O,  I_C  = '\x05', '\x06'  # italic
    # Сначала вложенные конструкции — единым маркером bold-italic
    h = re.sub(r'<strong>\s*<em>(.*?)</em>\s*</strong>',
               f'{BI_O}\\1{BI_C}', h, flags=re.DOTALL)
    h = re.sub(r'<em>\s*<strong>(.*?)</strong>\s*</em>',
               f'{BI_O}\\1{BI_C}', h, flags=re.DOTALL)
    # Потом одиночные
    h = re.sub(r'<strong>(.*?)</strong>', f'{B_O}\\1{B_C}', h, flags=re.DOTALL)
    h = re.sub(r'<b>(.*?)</b>',           f'{B_O}\\1{B_C}', h, flags=re.DOTALL)
    h = re.sub(r'<em>(.*?)</em>',         f'{I_O}\\1{I_C}', h, flags=re.DOTALL)
    h = re.sub(r'<i>(.*?)</i>',           f'{I_O}\\1{I_C}', h, flags=re.DOTALL)
    # Сдвиг пробела изнутри маркера наружу: "X<close>" вместо "X <close>"
    h = re.sub(rf'(\S) ([{B_C}{I_C}{BI_C}])', r'\1\2 ', h)
    h = re.sub(rf'([{B_O}{I_O}{BI_O}]) (\S)', r' \1\2', h)
    # И только теперь — плейсхолдеры в markdown-маркеры
    h = h.replace(BI_O, '***').replace(BI_C, '***')
    h = h.replace(B_O,  '**' ).replace(B_C,  '**' )
    h = h.replace(I_O,  '*'  ).replace(I_C,  '*'  )

    h = re.sub(r'<a href="([^"]*)">(.*?)</a>', r'[\2](\1)', h, flags=re.DOTALL)

    # Таблицы
    h = re.sub(r'<table[^>]*>(.*?)</table>', _table_match_to_md, h, flags=re.DOTALL)

    # Списки
    h = re.sub(r'<ul[^>]*>(.*?)</ul>', lambda m: _list_to_md(m.group(1), False), h, flags=re.DOTALL)
    h = re.sub(r'<ol[^>]*>(.*?)</ol>', lambda m: _list_to_md(m.group(1), True), h, flags=re.DOTALL)

    # Заголовки (с отступами для совместимости с convert.py split на \n\n)
    for level in range(1, 4):
        h = re.sub(
            f'<h{level}[^>]*>(.*?)</h{level}>',
            f'\n\n{"#" * level} \\1\n\n',
            h,
            flags=re.DOTALL,
        )

    # Параграфы → текст + двойной перенос
    h = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', h, flags=re.DOTALL)

    h = re.sub(r'<br\s*/?>', '\n', h)

    # Убираем оставшиеся HTML-теги
    h = re.sub(r'<[^>]+>', '', h)

    h = html_unescape(h)

    return h


def _table_match_to_md(match: re.Match) -> str:
    """Конвертирует HTML <table> в markdown-таблицу."""
    table_html = match.group(0)
    # Убираем thead/tbody обёртки
    table_html = re.sub(r'</?(?:thead|tbody|tfoot)[^>]*>', '', table_html)

    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL)
    if not rows:
        return ''

    md_rows = []
    for row_html in rows:
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL)
        cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
        md_rows.append('| ' + ' | '.join(cells) + ' |')

    if len(md_rows) >= 1:
        num_cols = len(re.findall(r'<t[dh][^>]*>', rows[0], re.DOTALL))
        separator = '| ' + ' | '.join(['---'] * max(num_cols, 1)) + ' |'
        md_rows.insert(1, separator)

    return '\n' + '\n'.join(md_rows) + '\n'


def _list_to_md(list_html: str, ordered: bool) -> str:
    """Конвертирует содержимое <ul>/<ol> в markdown-список."""
    items = re.findall(r'<li[^>]*>(.*?)</li>', list_html, re.DOTALL)
    lines = []
    for i, item in enumerate(items):
        item = re.sub(r'<[^>]+>', '', item).strip()
        if ordered:
            lines.append(f'{i + 1}. {item}')
        else:
            lines.append(f'- {item}')
    return '\n' + '\n'.join(lines) + '\n'


def _postprocess_md(md_text: str) -> str:
    """Нормализует markdown для совместимости с convert.py."""
    lines = [line.rstrip() for line in md_text.split('\n')]
    md_text = '\n'.join(lines)

    md_text = re.sub(r'\n{3,}', '\n\n', md_text)

    # Пустая строка перед заголовком
    md_text = re.sub(r'([^\n])\n(#{1,3} )', r'\1\n\n\2', md_text)

    # Пустая строка после заголовка
    md_text = re.sub(r'(#{1,3} [^\n]+)\n([^\n#])', r'\1\n\n\2', md_text)

    # Пустая строка после строки таблицы, если следом не таблица и не пустая строка
    md_text = re.sub(r'(\|[^\n]*\|)\n(?!\||\n)', r'\1\n\n', md_text)

    # Пустая строка после строки списка перед параграфом-не-списком
    md_text = re.sub(
        r'((?:^- |^\d+\. )[^\n]+)\n(?!- |\d+\. |\n)',
        r'\1\n\n',
        md_text,
        flags=re.MULTILINE,
    )

    return md_text.strip()


# --- Legacy: старая реализация на случай отката ---

def _docx_to_md_legacy(file_bytes: bytes) -> tuple[str, list]:
    """
    Legacy: Конвертирует DOCX в Markdown через python-docx.
    Сохранена как запасной вариант.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document(io.BytesIO(file_bytes))
    lines = []
    images = _extract_images_from_docx(doc)

    font_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)
    avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12.0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        style_name = para.style.name if para.style else ''

        heading_level = 0

        if 'Heading 1' in style_name or style_name == 'Title':
            heading_level = 1
        elif 'Heading 2' in style_name:
            heading_level = 2
        elif 'Heading 3' in style_name:
            heading_level = 3

        if heading_level == 0:
            run_sizes = [r.font.size.pt for r in para.runs if r.font.size]
            run_bolds = [r.bold for r in para.runs if r.text.strip()]
            is_bold = all(run_bolds) and len(run_bolds) > 0
            max_size = max(run_sizes) if run_sizes else avg_size

            if is_bold and max_size >= avg_size * 1.5:
                heading_level = 1
            elif is_bold and max_size >= avg_size * 1.25:
                heading_level = 2
            elif is_bold and max_size >= avg_size * 1.1:
                heading_level = 3

        if heading_level == 1:
            lines.append(f'# {text}')
        elif heading_level == 2:
            lines.append(f'## {text}')
        elif heading_level == 3:
            lines.append(f'### {text}')
        else:
            md_text = _runs_to_md_legacy(para.runs)
            if md_text.strip():
                lines.append(md_text)

        lines.append('')

    return '\n'.join(lines), images


def _runs_to_md_legacy(runs) -> str:
    """Legacy: переводит runs параграфа в MD с сохранением bold/italic."""
    result = ''
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            result += f'***{text}***'
        elif run.bold:
            result += f'**{text}**'
        elif run.italic:
            result += f'*{text}*'
        else:
            result += text
    return result


# =============================================================================
# PDF → MD
# =============================================================================

def pdf_to_md(file_bytes: bytes) -> tuple[str, list]:
    """
    Конвертирует PDF в Markdown.
    Заголовки определяются по размеру шрифта относительно среднего.
    Изображения извлекаются через pdfplumber (если возможно).
    """
    try:
        import pdfplumber
    except ImportError:
        return _pdf_fallback(file_bytes), []

    lines = []
    images = []
    image_counter = [0]

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        # Шаг 1: определяем средний размер шрифта по всему документу
        all_sizes = []
        for page in pdf.pages:
            for word in page.extract_words(extra_attrs=['size', 'fontname']):
                try:
                    all_sizes.append(float(word.get('size', 12)))
                except (ValueError, TypeError):
                    pass
        avg_size = sum(all_sizes) / len(all_sizes) if all_sizes else 12.0

        # Шаг 2: извлекаем текст постранично
        prev_text = ''
        for page_num, page in enumerate(pdf.pages):
            words = page.extract_words(
                extra_attrs=['size', 'fontname'],
                x_tolerance=3,
                y_tolerance=3
            )

            # Группируем слова в строки по Y-координате
            lines_on_page = {}
            for word in words:
                y_key = round(float(word.get('top', 0)), 0)
                if y_key not in lines_on_page:
                    lines_on_page[y_key] = []
                lines_on_page[y_key].append(word)

            for y_key in sorted(lines_on_page.keys()):
                line_words = lines_on_page[y_key]
                text = ' '.join(w['text'] for w in line_words).strip()
                if not text or text == prev_text:
                    continue
                prev_text = text

                # Определяем размер шрифта строки
                sizes = []
                for w in line_words:
                    try:
                        sizes.append(float(w.get('size', avg_size)))
                    except (ValueError, TypeError):
                        sizes.append(avg_size)
                line_size = max(sizes) if sizes else avg_size

                # Определяем жирность по имени шрифта
                fontnames = [w.get('fontname', '') for w in line_words]
                is_bold = any('Bold' in f or 'bold' in f or 'Heavy' in f
                              for f in fontnames)

                # Определяем уровень заголовка
                ratio = line_size / avg_size if avg_size > 0 else 1.0

                if ratio >= 1.5 or (ratio >= 1.3 and is_bold):
                    lines.append(f'# {text}')
                elif ratio >= 1.25 or (ratio >= 1.1 and is_bold):
                    lines.append(f'## {text}')
                elif ratio >= 1.1 and is_bold:
                    lines.append(f'### {text}')
                else:
                    lines.append(text)

                lines.append('')

            # Извлекаем изображения со страницы
            try:
                for img in page.images:
                    img_bytes = _extract_pdf_image(page, img)
                    if img_bytes:
                        fname = f"image_p{page_num}_{image_counter[0]}.png"
                        images.append((fname, img_bytes))
                        image_counter[0] += 1
            except Exception:
                pass

    return '\n'.join(lines), images


def _extract_pdf_image(page, img_dict) -> Optional[bytes]:
    """Пытается извлечь изображение из PDF страницы."""
    try:
        x0 = img_dict['x0']
        y0 = img_dict['y0']
        x1 = img_dict['x1']
        y1 = img_dict['y1']
        cropped = page.crop((x0, y0, x1, y1))
        img_obj = cropped.to_image(resolution=150)
        buf = io.BytesIO()
        img_obj.save(buf, format='PNG')
        return buf.getvalue()
    except Exception:
        return None


def _pdf_fallback(file_bytes: bytes) -> str:
    """Запасной вариант если pdfplumber не установлен."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        lines = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
                lines.append('')
        return '\n'.join(lines)
    except Exception:
        return "❌ Не удалось извлечь текст из PDF. Установите pdfplumber."


# =============================================================================
# TXT → MD
# =============================================================================

def txt_to_md(file_bytes: bytes) -> tuple[str, list]:
    """
    Конвертирует TXT в Markdown.
    В TXT нет форматирования — просто разбиваем на абзацы.
    Пустые строки = разделители абзацев.
    """
    try:
        text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        text = file_bytes.decode('cp1251', errors='replace')

    # Нормализуем переносы строк
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Разбиваем на абзацы по двойному переносу
    paragraphs = re.split(r'\n{2,}', text)
    lines = []
    for para in paragraphs:
        para = para.strip()
        if para:
            # Объединяем одиночные переносы внутри абзаца в пробел
            para = re.sub(r'\n', ' ', para)
            lines.append(para)
            lines.append('')

    return '\n'.join(lines), []


# =============================================================================
# УНИВЕРСАЛЬНЫЙ ВХОД
# =============================================================================

def convert_file_to_md(file_bytes: bytes, filename: str) -> tuple[str, list]:
    """
    Главная функция — определяет формат по расширению и вызывает нужный конвертер.
    Возвращает (md_text, images).
    """
    ext = filename.lower().split('.')[-1]

    if ext == 'docx':
        return docx_to_md(file_bytes)
    elif ext == 'pdf':
        return pdf_to_md(file_bytes)
    elif ext in ('txt', 'text'):
        return txt_to_md(file_bytes)
    else:
        raise ValueError(f"Неподдерживаемый формат: .{ext}")
