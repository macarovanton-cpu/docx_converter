import io
import os
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# =============================================================================
# НАСТРОЙКИ — МЕНЯЙ ТОЛЬКО ЭТИ ТРИ СТРОКИ
# =============================================================================
INPUT_FILE    = r"C:\Users\tonik\Desktop\docx_converter\Cheremhovo.md"       # ← твой .md файл
OUTPUT_FILE   = r"C:\Users\tonik\Desktop\docx_converter\Cheremhovo.docx"     # ← куда сохранить
TEMPLATE_FILE = r"C:\Users\tonik\Desktop\docx_converter\template.docx"   # ← шаблон с хедером
# =============================================================================

# === ЦВЕТОВАЯ ПАЛИТРА ===
BRAND_BLUE      = "015198"
BRAND_RED       = "D04514"
BRAND_ORANGE    = "EF7F1A"
BRAND_WHITE     = "FFFFFF"
TEXT_DARK       = "1A1A1A"
BG_LIGHT_BLUE   = "EBF3FB"
BG_LIGHT_ORANGE = "FFF8F0"
BG_TABLE_ROW    = "F0F5FA"
BORDER_LIGHT    = "CCCCCC"
TEXT_MUTED      = "888888"
COLOR_YES       = "1E7A34"   # зелёный для ✓
COLOR_NO        = "C0392B"   # красный для ✗

# ПРАВКА #17: авто-замена ячеек «Да» → «✓ Да», «Нет»/«Отсутствует» → «✗ ...»
# Установи False, если хочешь сохранять текст ячеек как есть.
ENABLE_TABLE_SYMBOLS = True

# Поля шаблона: left=2cm, right=1.5cm → рабочая ширина 17.5cm
CONTENT_WIDTH_CM = 17.5


# =============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =============================================================================

def clear_body(doc):
    """Удаляет всё содержимое тела, сохраняя финальный sectPr."""
    body = doc.element.body
    to_remove = [c for c in body
                 if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) != 'sectPr']
    for el in to_remove:
        body.remove(el)
    body.insert(0, OxmlElement('w:p'))


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_margins_and_borders(cell, hex_color, sz):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, w in [('top','100'),('bottom','100'),('left','160'),('right','160')]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), w)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), str(sz))
        bdr.set(qn('w:color'), hex_color)
        tcBorders.append(bdr)
    tcPr.append(tcBorders)


def set_cell_no_borders(cell):
    """Убирает все видимые границы ячейки (для таблиц-обёрток)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'none')
        bdr.set(qn('w:sz'), '0')
        bdr.set(qn('w:color'), 'auto')
        tcBorders.append(bdr)
    tcPr.append(tcBorders)


def add_paragraph_border(paragraph, side, color, size, space="0"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), str(space))
    border.set(qn('w:color'), color)
    pBdr.append(border)


def add_paragraph_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def set_table_width_dxa(table, width_cm):
    """Фиксирует ширину таблицы. 1cm = 567 DXA."""
    dxa = int(width_cm * 567)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_table_no_spacing(table):
    """Убирает межтабличные отступы (для таблиц-обёрток)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'), '0')
    spacing.set(qn('w:type'), 'dxa')
    tblPr.append(spacing)


def set_row_height(row, height_dxa):
    """Устанавливает минимальную высоту строки таблицы."""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        row._tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_dxa))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def set_keep_with_next(paragraph):
    """Параграф остаётся на той же странице что и следующий."""
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    pPr.append(kwn)


def set_keep_together(paragraph):
    """Не разрывает параграф по страницам."""
    pPr = paragraph._p.get_or_add_pPr()
    kt = OxmlElement('w:keepLines')
    pPr.append(kt)


def set_run_font(run, name, size_pt, color_hex, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


# ПРАВКА #21: обработка markdown-ссылок [text](url)
def add_hyperlink_run(paragraph, url, text, font_name='PT Sans', font_size=12,
                      bold=False, italic=False):
    if not url or not url.strip():
        return None
    url = url.strip()
    url = re.sub(r'\\([._\-+~#?&=/])', r'\1', url)      # ПРАВКА #24: раскрытие markdown-экранирования в URL
    if not re.match(r'^https?://', url):
        url = 'https://' + url
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), BRAND_BLUE)
    rPr.append(color_el)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size * 2))
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    run_el.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number_field(run):
    for ftype in ['begin', None, 'separate', 'end']:
        if ftype is None:
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = 'PAGE'
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), ftype)
        run._r.append(el)


# =============================================================================
# ДЕТЕКТОРЫ ТИПОВ БЛОКОВ
# =============================================================================

def is_stage_paragraph(text):
    return bool(re.match(r'^\*\*(Стадия|Фаза|Шаг|Этап|ВАЖНО)', text, re.IGNORECASE))

def is_photo_placeholder(text):
    return '📷' in text or '[Место для фото' in text

# ПРАВКА #23: блок-картинка ![alt](src)
_IMG_BLOCK_RE = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')

def is_requisites_block(text):
    return bool(re.match(r'^\*\*(Кому|От кого|Кому:|От кого:)', text))

def is_signature_block(text):
    return bool(re.match(r'^\*?С уважением', text))

def is_callout_block(text):
    """Блок !! текст !! — callout-врезка."""
    return text.startswith('!!') and text.endswith('!!')


# =============================================================================
# INLINE MARKDOWN ПАРСЕР
# =============================================================================

def _parse_bold_italic(paragraph, text, font_name, font_size,
                       font_color, is_italic_base):
    # ПРАВКА #22: поддержка ***bold-italic*** в inline-парсере
    pattern = re.compile(r'(\*\*\*[^*\n]+?\*\*\*|\*\*[^*\n]+?\*\*|\*(?!\*)[^*\n]+?\*(?!\*))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        is_bold   = False
        is_italic = is_italic_base
        clean_text = part
        if part.startswith('***') and part.endswith('***') and len(part) >= 7:
            clean_text = part[3:-3]
            is_bold = True
            is_italic = True
        elif part.startswith('**') and part.endswith('**') and len(part) >= 5:
            clean_text = part[2:-2]
            is_bold = True
        elif (part.startswith('*') and part.endswith('*')
              and len(part) >= 3 and not part.startswith('**')):
            clean_text = part[1:-1]
            is_italic = True
        set_run_font(run, font_name, font_size, font_color,
                     bold=is_bold, italic=is_italic)
        run.text = clean_text


def parse_inline_markdown(paragraph, text, font_name='PT Sans', font_size=12,
                          font_color=TEXT_DARK, is_italic_base=False,
                          images=None, content_width_cm=None):
    """Обрабатывает ***жирный-курсив***, **жирный**, *курсив*, [ссылки](url)
    и инлайн-картинки ![alt](src) (ПРАВКА #32)."""
    text = re.sub(r'\\([.\-+_)(:!=])', r'\1', text)      # ПРАВКА #24: раскрытие markdown-экранирования \X → X
    # ПРАВКА #21: сначала разбиваем по ссылкам [text](url)
    # ПРАВКА #32: ![alt](src) распознаётся ДО ссылок — раньше «!» оставался
    # литералом, а src превращался в мусорную гиперссылку https://image_1.png
    link_re = re.compile(r'(!?\[[^\]]*?\]\([^)]*?\))')
    img_detail = re.compile(r'^!\[([^\]]*?)\]\(([^)]*?)\)$')
    link_detail = re.compile(r'^\[([^\]]+?)\]\(([^)]*?)\)$')
    for segment in link_re.split(text):
        if not segment:
            continue
        mi = img_detail.match(segment)
        if mi:
            alt, src = mi.group(1).strip(), mi.group(2).strip()
            if images and src in images:
                run = paragraph.add_run()
                run.add_picture(io.BytesIO(images[src]),
                                width=_image_width(images[src], content_width_cm))
            elif alt:
                _parse_bold_italic(paragraph, alt, font_name,
                                   font_size, font_color, is_italic_base)
            continue
        m = link_detail.match(segment)
        if m:
            link_text, link_url = m.group(1), m.group(2).strip()
            if link_url:
                add_hyperlink_run(paragraph, link_url, link_text,
                                  font_name, font_size)
            else:
                _parse_bold_italic(paragraph, link_text, font_name,
                                   font_size, font_color, is_italic_base)
        else:
            _parse_bold_italic(paragraph, segment, font_name,
                               font_size, font_color, is_italic_base)


# ПРАВКА #32: расчёт ширины картинки вынесен из _add_inline_image,
# используется и блочной, и инлайн-вставкой
def _image_width(img_bytes, content_width_cm):
    if content_width_cm is None:
        content_width_cm = CONTENT_WIDTH_CM
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        w_px, _ = img.size
        img.close()
        w_cm = w_px / 96 * 2.54
        return Cm(min(w_cm, content_width_cm))
    except Exception:
        return Cm(content_width_cm)


def _add_inline_image(doc, img_bytes, content_width_cm):
    import tempfile
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    tmp.write(img_bytes)
    tmp.close()
    width = _image_width(img_bytes, content_width_cm)   # ПРАВКА #32
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(tmp.name, width=width)
    os.unlink(tmp.name)


# =============================================================================
# СПЕЦИАЛЬНЫЕ БЛОКИ-КОНСТРУКТОРЫ
# =============================================================================

def add_intro_paragraph(doc, block, content_width_cm):
    """
    ПРАВКА #4: Вводный абзац после H1 — таблица-обёртка с цветной левой полосой.
    Выглядит как акцентный callout для главной мысли документа.
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = table.allow_autofit = False
    set_table_width_dxa(table, content_width_cm)
    set_table_no_spacing(table)

    cell = table.rows[0].cells[0]

    # Только левая граница — фирменный синий, 2pt
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'none')
        bdr.set(qn('w:sz'), '0')
        bdr.set(qn('w:color'), 'auto')
        tcBorders.append(bdr)
    left_bdr = OxmlElement('w:left')
    left_bdr.set(qn('w:val'), 'single')
    left_bdr.set(qn('w:sz'), '18')   # ~2.25pt
    left_bdr.set(qn('w:space'), '4')
    left_bdr.set(qn('w:color'), BRAND_BLUE)
    tcBorders.append(left_bdr)
    tcPr.append(tcBorders)

    # Внутренние отступы ячейки
    tcMar = OxmlElement('w:tcMar')
    for side, w in [('top','80'),('bottom','80'),('left','220'),('right','0')]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), w)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.alignment  = WD_ALIGN_PARAGRAPH.LEFT  # ПРАВКА #27
    p.paragraph_format.space_after = Pt(0)
    # ПРАВКА #1: межстрочный 1.3
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.3

    lines = block.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        if not line: continue
        if i > 0: p.add_run().add_break()
        parse_inline_markdown(p, line)


def add_callout_box(doc, text, content_width_cm):
    """
    ПРАВКА #6: Callout-врезка !! текст !! — таблица с заливкой и бордером.
    Используется для формул, ключевых выводов, важных цифр.
    """
    clean = text.strip('!').strip()
    table = doc.add_table(rows=1, cols=1)
    table.autofit = table.allow_autofit = False
    set_table_width_dxa(table, content_width_cm)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F2F6FA')   # очень лёгкий голубой

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')     # 0.5pt
        bdr.set(qn('w:color'), 'C5D8EC')
        tcBorders.append(bdr)
    # Акцент — левая граница чуть толще
    tcPr.append(tcBorders)

    tcMar = OxmlElement('w:tcMar')
    for side, w in [('top','140'),('bottom','140'),('left','220'),('right','220')]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), w)
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.alignment  = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    parse_inline_markdown(p, clean, 'PT Sans', 12, BRAND_BLUE)
    for r in p.runs:
        r.bold = True


# ПРАВКА #30: компактный пустой параграф-спейсер после таблиц-блоков,
# чтобы соседние w:tbl не склеивались Word'ом в одну таблицу
def add_compact_spacer(doc):
    sp = doc.add_paragraph()
    pPr = sp._p.get_or_add_pPr()
    s = OxmlElement('w:spacing')
    s.set(qn('w:before'), '0')
    s.set(qn('w:after'), '120')
    s.set(qn('w:line'), '120')
    s.set(qn('w:lineRule'), 'exact')
    pPr.append(s)


def add_table_cell_content(p, text, font_size=10):
    """
    ПРАВКА #8: Добавляет ✓/✗ перед значениями «Да»/«Нет» в ячейках таблицы.
    Применяется для любой сравнительной таблицы автоматически.
    ПРАВКА #17: поведение управляется флагом ENABLE_TABLE_SYMBOLS.
    """
    stripped = text.strip()

    if not ENABLE_TABLE_SYMBOLS:
        parse_inline_markdown(p, stripped, 'PT Sans', font_size, TEXT_DARK)
        return

    # Проверяем начало ячейки на Да/Нет
    if re.match(r'^Да\b', stripped, re.IGNORECASE):
        icon_run = p.add_run('✓ ')
        set_run_font(icon_run, 'PT Sans', font_size, COLOR_YES, bold=True)
        rest = re.sub(r'^Да\b\s*', '', stripped, flags=re.IGNORECASE)
        if rest:
            parse_inline_markdown(p, rest, 'PT Sans', font_size, TEXT_DARK)
        else:
            run = p.add_run('Да')
            set_run_font(run, 'PT Sans', font_size, TEXT_DARK)
    elif re.match(r'^Нет\b', stripped, re.IGNORECASE):
        icon_run = p.add_run('✗ ')
        set_run_font(icon_run, 'PT Sans', font_size, COLOR_NO, bold=True)
        rest = re.sub(r'^Нет\b\s*', '', stripped, flags=re.IGNORECASE)
        if rest:
            parse_inline_markdown(p, rest, 'PT Sans', font_size, TEXT_DARK)
        else:
            run = p.add_run('Нет')
            set_run_font(run, 'PT Sans', font_size, TEXT_DARK)
    elif re.match(r'^Отсутствует\b', stripped, re.IGNORECASE):
        icon_run = p.add_run('✗ ')
        set_run_font(icon_run, 'PT Sans', font_size, COLOR_NO, bold=True)
        rest = re.sub(r'^Отсутствует\b\s*', '', stripped, flags=re.IGNORECASE)
        parse_inline_markdown(p, ('Отсутствует ' + rest).strip(), 'PT Sans', font_size, TEXT_DARK)
    else:
        parse_inline_markdown(p, stripped, 'PT Sans', font_size, TEXT_DARK)


# =============================================================================
# ПРАВКА #18: АВТОМАТИЧЕСКИЕ ПЕРЕНОСЫ СЛОВ
# =============================================================================

def enable_auto_hyphenation(doc):
    """
    ПРАВКА #18: включает автоматические переносы слов на уровне документа.
    Без этого justify создаёт большие пробелы между словами в коротких строках.
    """
    settings = doc.settings.element
    existing = settings.find(qn('w:autoHyphenation'))
    if existing is not None:
        settings.remove(existing)
    auto_hyphen = OxmlElement('w:autoHyphenation')
    settings.insert(0, auto_hyphen)
    do_not_hyphen_caps = OxmlElement('w:doNotHyphenateCaps')
    settings.insert(1, do_not_hyphen_caps)


# =============================================================================
# ПРАВКА #13: НУМЕРАЦИЯ СПИСКОВ
# =============================================================================

def ensure_list_numbering(doc):
    """
    ПРАВКА #13: гарантирует наличие в numbering.xml кастомных определений
    для bullet (•) и numbered (1. 2. 3.) списков.
    Возвращает (bullet_num_id, numbered_num_id).
    """

    BULLET_ABSTRACT_ID = 100
    NUMBERED_ABSTRACT_ID = 101
    BULLET_NUM_ID = 100
    NUMBERED_NUM_ID = 101

    # Получаем или создаём numbering part
    try:
        numbering_part = doc.part.numbering_part
    except (KeyError, AttributeError):
        # Нет numbering part — создаём заглушку и регистрируем
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree

        numbering_uri = PackURI('/word/numbering.xml')
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        numbering_xml = etree.Element(qn('w:numbering'), nsmap=nsmap)
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml'
        part = Part(
            numbering_uri,
            content_type,
            etree.tostring(numbering_xml, xml_declaration=True, encoding='UTF-8', standalone=True),
            doc.part.package
        )
        doc.part.relate_to(part, RT.NUMBERING)
        numbering_part = doc.part.numbering_part

    numbering_elem = numbering_part.element

    # Проверяем, не созданы ли уже наши определения
    existing_abstract = {
        int(a.get(qn('w:abstractNumId')))
        for a in numbering_elem.findall(qn('w:abstractNum'))
    }
    if BULLET_ABSTRACT_ID in existing_abstract and NUMBERED_ABSTRACT_ID in existing_abstract:
        return BULLET_NUM_ID, NUMBERED_NUM_ID

    # --- Bullet abstractNum ---
    if BULLET_ABSTRACT_ID not in existing_abstract:
        abstract_bullet = OxmlElement('w:abstractNum')
        abstract_bullet.set(qn('w:abstractNumId'), str(BULLET_ABSTRACT_ID))
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'bullet')
        lvl.append(numFmt)
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '•')
        lvl.append(lvlText)
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Symbol')
        rFonts.set(qn('w:hAnsi'), 'Symbol')
        rFonts.set(qn('w:hint'), 'default')
        rPr.append(rFonts)
        lvl.append(rPr)
        abstract_bullet.append(lvl)
        numbering_elem.append(abstract_bullet)

    # --- Numbered abstractNum ---
    if NUMBERED_ABSTRACT_ID not in existing_abstract:
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(NUMBERED_ABSTRACT_ID))
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'decimal')
        lvl.append(numFmt)
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '%1.')
        lvl.append(lvlText)
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)
        abstract_num.append(lvl)
        numbering_elem.append(abstract_num)

    # --- num элементы (ссылки на abstractNum) ---
    existing_num = {
        int(n.get(qn('w:numId')))
        for n in numbering_elem.findall(qn('w:num'))
    }
    if BULLET_NUM_ID not in existing_num:
        num_bullet = OxmlElement('w:num')
        num_bullet.set(qn('w:numId'), str(BULLET_NUM_ID))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(BULLET_ABSTRACT_ID))
        num_bullet.append(abstract_ref)
        numbering_elem.append(num_bullet)

    if NUMBERED_NUM_ID not in existing_num:
        num_numbered = OxmlElement('w:num')
        num_numbered.set(qn('w:numId'), str(NUMBERED_NUM_ID))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(NUMBERED_ABSTRACT_ID))
        num_numbered.append(abstract_ref)
        numbering_elem.append(num_numbered)

    return BULLET_NUM_ID, NUMBERED_NUM_ID


def new_numbered_num_id(doc):
    """
    ПРАВКА #31: свежий w:num со startOverride=1 для каждого нового списка.
    Без этого все нумерованные списки документа делили один счётчик и
    второй список продолжался с 4 вместо 1.
    """
    numbering_elem = doc.part.numbering_part.element
    existing = [int(n.get(qn('w:numId')))
                for n in numbering_elem.findall(qn('w:num'))]
    num_id = max(existing) + 1
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), '101')   # NUMBERED_ABSTRACT_ID из ensure_list_numbering
    num.append(abstract_ref)
    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    override.append(start_override)
    num.append(override)
    numbering_elem.append(num)
    return num_id


def set_paragraph_numbering(paragraph, num_id, ilvl=0):
    """ПРАВКА #13: привязывает абзац к numbering definition через OXML."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


# =============================================================================
# ОСНОВНАЯ ЛОГИКА КОНВЕРТАЦИИ
# =============================================================================

def convert_md_to_docx(md_text, output_filename, template_path=None, images=None):

    # ПРАВКА #28: нормализация переводов строк — CRLF/CR ломали split('\n\n') и regex #26
    md_text = md_text.replace('\r\n', '\n').replace('\r', '\n')

    # --- Открываем шаблон или создаём чистый документ ---
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        clear_body(doc)
        content_width_cm = CONTENT_WIDTH_CM
        print(f"  Шаблон: {template_path}")
    else:
        doc = Document()
        section = doc.sections[0]
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        content_width_cm = 21.0 - 2.54 * 2
        print("  ⚠️  Шаблон не найден — хедер не будет добавлен")

        footer   = doc.sections[0].footer
        footer_p = footer.paragraphs[0]
        footer_p.paragraph_format.tab_stops.add_tab_stop(
            Cm(content_width_cm), WD_TAB_ALIGNMENT.RIGHT)
        rl = footer_p.add_run("ООО «ТПК «Тензосила»")
        set_run_font(rl, 'PT Sans', 9, TEXT_MUTED)
        footer_p.add_run("\t")
        rr = footer_p.add_run()
        set_run_font(rr, 'PT Sans', 9, TEXT_MUTED)
        add_page_number_field(rr)

    # --- Базовый стиль Normal ---
    sn = doc.styles['Normal']
    sn.font.name      = 'PT Sans'
    sn.font.size      = Pt(12)
    sn.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    sn.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT  # ПРАВКА #27: justify → left для читаемости с латинской терминологией
    sn.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sn.paragraph_format.line_spacing      = 1.3   # ПРАВКА #1: было 1.2
    sn.paragraph_format.space_after       = Pt(8)  # ПРАВКА #2: было Pt(6)

    # ПРАВКА #18: включаем автоматические переносы слов
    enable_auto_hyphenation(doc)

    # ПРАВКА #13: регистрируем numbering для bullet/numbered списков
    bullet_num_id, numbered_num_id = ensure_list_numbering(doc)

    # ПРАВКА #23: позиционирование картинок по тексту
    _images_dict = {}
    if images:
        for fname, img_bytes in images:
            _images_dict[fname] = img_bytes

    # ПРАВКА #26: bold prefix + hard break → отдельные параграфы
    # ПРАВКА #29: не резать блоки реквизитов/стадий — их префиксы ловят
    # is_requisites_block (case-sensitive) и is_stage_paragraph (IGNORECASE)
    md_text = re.sub(
        r'^(\*\*(?!(?:Кому|От кого|(?i:Стадия|Фаза|Шаг|Этап|ВАЖНО)))[^*\n]{1,100}?:\*\*)  +\n(?!\n)',
        r'\1\n\n',
        md_text,
        flags=re.MULTILINE,
    )

    # --- Парсинг блоков Markdown ---
    blocks = md_text.split('\n\n')

    after_heading         = False
    # ПРАВКА #12: единый флаг — intro-блок только сразу после H1
    pending_intro_after_h1 = False
    last_list_paragraph   = None
    last_regular_paragraph = None  # ПРАВКА #10: для keep_with_next перед подписью
    current_numbered_num_id = None  # ПРАВКА #31: numId текущего нумерованного списка

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        is_list_item = (block.startswith('- ')
                        or block.startswith('* ')
                        # ПРАВКА #31: пробел после точки обязателен, маркер — max 2
                        # цифры, иначе «2025. Год…» съедался как пункт списка
                        or bool(re.match(r'^\d{1,2}\. ', block)))
        if not is_list_item:
            # ПРАВКА #31: любой не-списочный блок завершает текущий
            # нумерованный список — следующий начнётся с 1
            current_numbered_num_id = None
        if not is_list_item and last_list_paragraph:
            last_list_paragraph.paragraph_format.space_after = Pt(10)
            last_list_paragraph = None

        # ── H1 ───────────────────────────────────────────────────────────────
        if block.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after  = Pt(12)
            parse_inline_markdown(p, block[2:], 'PT Sans Narrow', 18, BRAND_BLUE)
            for r in p.runs: r.bold = True
            # ПРАВКА #14: H1 не отрывается от контента ниже
            set_keep_with_next(p)
            dec = doc.add_paragraph()
            dec.paragraph_format.space_before = Pt(0)
            dec.paragraph_format.space_after  = Pt(10)
            add_paragraph_border(dec, 'bottom', BRAND_RED, 12)
            # ПРАВКА #14: декоративная линия тоже держится с контентом ниже
            set_keep_with_next(dec)
            after_heading = True
            # ПРАВКА #12: intro-блок ожидается только сразу после H1
            pending_intro_after_h1 = True

        # ── H2 ───────────────────────────────────────────────────────────────
        elif block.startswith('## '):
            pending_intro_after_h1 = False   # ПРАВКА #12
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
            # ПРАВКА #3: H2 остаётся цветным 14pt
            parse_inline_markdown(p, block[3:], 'PT Sans Narrow', 14, BRAND_RED)
            for r in p.runs: r.bold = True
            # ПРАВКА #14: H2 не отрывается от контента ниже
            set_keep_with_next(p)
            after_heading = True

        # ── H3 ───────────────────────────────────────────────────────────────
        elif block.startswith('### '):
            pending_intro_after_h1 = False   # ПРАВКА #12
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            # ПРАВКА #3: H3 — чёрный bold (отличается от H2 цветом и размером)
            # ПРАВКА #15: размер 11pt → 13pt, чтобы H3 был крупнее тела (12pt)
            parse_inline_markdown(p, block[4:], 'PT Sans Narrow', 13, TEXT_DARK)
            for r in p.runs: r.bold = True
            # ПРАВКА #14: H3 не отрывается от контента ниже
            set_keep_with_next(p)
            after_heading = True

        # ── Цитаты > ─────────────────────────────────────────────────────────
        elif block.startswith('>'):
            pending_intro_after_h1 = False   # ПРАВКА #12
            clean = '\n'.join([l.lstrip('> ') for l in block.split('\n')])
            p = doc.add_paragraph()
            # ПРАВКА #5: увеличен отступ, лёгкая заливка
            p.paragraph_format.left_indent  = Cm(1.8)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(10)
            add_paragraph_border(p, 'left', BRAND_ORANGE, 18, space=4)
            add_paragraph_shading(p, 'F7F7F7')
            parse_inline_markdown(p, clean, 'PT Sans', 12, "555555", is_italic_base=True)
            last_regular_paragraph = p
            after_heading = False

        # ── Callout-врезка !! текст !! ────────────────────────────────────────
        elif is_callout_block(block):
            pending_intro_after_h1 = False   # ПРАВКА #12
            # ПРАВКА #6: новый тип блока — оформляется как акцентная таблица
            add_callout_box(doc, block, content_width_cm)
            # ПРАВКА #30: спейсер, иначе callout склеивается со следующим w:tbl
            add_compact_spacer(doc)
            after_heading = False

        # ── Встроенные картинки (ПРАВКА #23) ─────────────────────────────────
        # ПРАВКА #32: ветка работает и при пустом images — раньше блок падал
        # в inline-парсер и превращался в «!» + мусорную гиперссылку
        elif _IMG_BLOCK_RE.match(block):
            pending_intro_after_h1 = False
            m = _IMG_BLOCK_RE.match(block)
            alt, img_src = m.group(1).strip(), m.group(2).strip()
            if img_src in _images_dict:
                _add_inline_image(doc, _images_dict[img_src], content_width_cm)
            elif is_photo_placeholder(alt):
                # alt — фото-плейсхолдер: рендерим как ветку 📷 ниже
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(1.0)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(10)
                add_paragraph_border(p, 'left', BRAND_ORANGE, 18)
                add_paragraph_shading(p, BG_LIGHT_ORANGE)
                parse_inline_markdown(p, alt, 'PT Sans', 11, "999999",
                                      is_italic_base=True)
            else:
                p = doc.add_paragraph()
                parse_inline_markdown(
                    p, f'{alt} (изображение не найдено: {img_src})'.strip())
            after_heading = False

        # ── Плейсхолдеры фото ────────────────────────────────────────────────
        elif is_photo_placeholder(block):
            pending_intro_after_h1 = False   # ПРАВКА #12
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(10)
            add_paragraph_border(p, 'left', BRAND_ORANGE, 18)
            add_paragraph_shading(p, BG_LIGHT_ORANGE)
            parse_inline_markdown(p, block.replace('>', '').strip(),
                                  'PT Sans', 11, "999999", is_italic_base=True)
            after_heading = False

        # ── Стадии / ВАЖНО ───────────────────────────────────────────────────
        elif is_stage_paragraph(block):
            pending_intro_after_h1 = False   # ПРАВКА #12
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.75)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(6)
            add_paragraph_border(p, 'left', BRAND_BLUE, 12)
            add_paragraph_shading(p, BG_LIGHT_BLUE)
            parse_inline_markdown(p, block)
            after_heading = False

        # ── Списки ───────────────────────────────────────────────────────────
        elif is_list_item:
            pending_intro_after_h1 = False   # ПРАВКА #12
            for line in block.split('\n'):
                line = line.strip()
                if not line: continue
                is_num = bool(re.match(r'^\d{1,2}\. ', line))  # ПРАВКА #31: пробел обязателен, max 2 цифры
                # ПРАВКА #13: привязываем numbering напрямую через OXML
                # ПРАВКА #31: каждый новый нумерованный список — свой numId с рестартом
                if is_num and current_numbered_num_id is None:
                    current_numbered_num_id = new_numbered_num_id(doc)
                p = doc.add_paragraph()
                set_paragraph_numbering(p, current_numbered_num_id if is_num else bullet_num_id)
                p.paragraph_format.left_indent       = Cm(1.5)
                p.paragraph_format.first_line_indent = Cm(-0.75)
                p.paragraph_format.space_after       = Pt(4)
                parse_inline_markdown(p, re.sub(r'^(\- |\* |\d{1,2}\. )', '', line),  # ПРАВКА #31: синхронно с is_num
                                      images=_images_dict, content_width_cm=content_width_cm)  # ПРАВКА #32
                last_list_paragraph = p
                last_regular_paragraph = p
            after_heading = False

        # ── Таблицы ──────────────────────────────────────────────────────────
        elif block.startswith('|') and '\n|' in block:
            pending_intro_after_h1 = False   # ПРАВКА #12
            lines = [l.strip() for l in block.split('\n')
                     if l.strip() and not re.match(r'^\|[-|: ]+\|$', l.strip())]
            if not lines: continue

            headers  = [c.strip() for c in lines[0].strip('|').split('|')]
            n_cols   = len(headers)
            table    = doc.add_table(rows=1, cols=n_cols)
            # ПРАВКА #16: autofit для распределения ширин по содержимому
            table.autofit = True
            table.allow_autofit = True
            set_table_width_dxa(table, content_width_cm)
            # tblLayout=auto чтобы Word распределял ширины колонок
            tblPr = table._tbl.find(qn('w:tblPr'))
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'autofit')
            tblPr.append(tblLayout)

            # ПРАВКА #7: заголовочная строка — увеличена высота и шрифт
            set_row_height(table.rows[0], 560)   # ~1cm минимальная высота

            for i, h in enumerate(headers):
                if i < len(table.rows[0].cells):
                    cell = table.rows[0].cells[i]
                    set_cell_shading(cell, BRAND_BLUE)
                    set_cell_margins_and_borders(cell, BORDER_LIGHT, 4)
                    # ПРАВКА #19: минимальная ширина первой колонки для длинных подписей
                    if n_cols >= 3 and i == 0:
                        tcPr_w = cell._tc.get_or_add_tcPr()
                        existing_w = tcPr_w.find(qn('w:tcW'))
                        if existing_w is not None:
                            tcPr_w.remove(existing_w)
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(6.0 * 567)))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr_w.append(tcW)
                    p = cell.paragraphs[0]
                    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    # ПРАВКА #7: шрифт 10pt → 11pt в заголовке таблицы
                    parse_inline_markdown(p, h, 'PT Sans Narrow', 11, BRAND_WHITE)
                    for r in p.runs: r.bold = True

            for row_idx, line in enumerate(lines[1:]):
                cols      = [c.strip() for c in line.strip('|').split('|')]
                row_cells = table.add_row().cells
                is_even   = (row_idx % 2 == 1)

                for i, c in enumerate(cols):
                    if i < len(row_cells):
                        cell = row_cells[i]

                        # ПРАВКА #9: последняя колонка 3-колоночной таблицы
                        # получает фирменный голубой фон (выделяем «наш» столбец)
                        if n_cols == 3 and i == n_cols - 1:
                            bg = BG_LIGHT_BLUE
                        else:
                            bg = BG_TABLE_ROW if is_even else BRAND_WHITE

                        set_cell_shading(cell, bg)
                        set_cell_margins_and_borders(cell, BORDER_LIGHT, 4)
                        # ПРАВКА #19: минимальная ширина первой колонки для длинных подписей
                        if n_cols >= 3 and i == 0:
                            tcPr_w = cell._tc.get_or_add_tcPr()
                            existing_w = tcPr_w.find(qn('w:tcW'))
                            if existing_w is not None:
                                tcPr_w.remove(existing_w)
                            tcW = OxmlElement('w:tcW')
                            tcW.set(qn('w:w'), str(int(6.0 * 567)))
                            tcW.set(qn('w:type'), 'dxa')
                            tcPr_w.append(tcW)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(0)
                        # ПРАВКА #8: автоматические ✓/✗ для Да/Нет/Отсутствует
                        add_table_cell_content(p, c, font_size=10)

            # Компактный отступ после таблицы
            sp = doc.add_paragraph()
            pPr = sp._p.get_or_add_pPr()
            s = OxmlElement('w:spacing')
            s.set(qn('w:before'), '0')
            s.set(qn('w:after'), '120')
            s.set(qn('w:line'), '120')
            s.set(qn('w:lineRule'), 'exact')
            pPr.append(s)
            after_heading = False

        # ── Разделители --- ───────────────────────────────────────────────────
        elif block.startswith('---'):
            pending_intro_after_h1 = False   # ПРАВКА #12
            continue

        # ── Реквизиты «Кому / От кого» ───────────────────────────────────────
        elif is_requisites_block(block):
            pending_intro_after_h1 = False   # ПРАВКА #12
            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(14)
            p.paragraph_format.left_indent  = Cm(0.4)
            p.paragraph_format.right_indent = Cm(0.4)
            add_paragraph_shading(p, BG_LIGHT_BLUE)
            for i, line in enumerate(block.split('\n')):
                line = line.strip()
                if not line: continue
                if i > 0: p.add_run().add_break()
                parse_inline_markdown(p, line)
            last_regular_paragraph = p
            after_heading = False

        # ── Подпись «С уважением» ─────────────────────────────────────────────
        elif is_signature_block(block):
            pending_intro_after_h1 = False   # ПРАВКА #12
            # ПРАВКА #10: последний абзац перед подписью держим вместе с ней
            if last_regular_paragraph is not None:
                set_keep_with_next(last_regular_paragraph)

            p = doc.add_paragraph()
            p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(28)
            p.paragraph_format.space_after  = Pt(0)
            add_paragraph_border(p, 'top', BRAND_RED, 4, space=8)
            # ПРАВКА #11: весь блок подписи не разрывается по страницам
            set_keep_together(p)
            for i, line in enumerate(block.split('\n')):
                line = line.strip()
                if not line: continue
                if i > 0: p.add_run().add_break()
                parse_inline_markdown(p, line)
            after_heading = False

        # ── Обычные абзацы ────────────────────────────────────────────────────
        else:
            # ПРАВКА #4 + ПРАВКА #12: intro-блок только сразу после H1
            if pending_intro_after_h1:
                add_intro_paragraph(doc, block, content_width_cm)
                pending_intro_after_h1 = False
                after_heading = False
                # Добавляем пустой параграф-отступ после врезки
                sp = doc.add_paragraph()
                pPr = sp._p.get_or_add_pPr()
                s = OxmlElement('w:spacing')
                s.set(qn('w:before'), '0')
                s.set(qn('w:after'), '120')
                s.set(qn('w:line'), '120')
                s.set(qn('w:lineRule'), 'exact')
                pPr.append(s)
                continue

            p = doc.add_paragraph()
            if not after_heading:
                p.paragraph_format.first_line_indent = Cm(0.75)
            for i, line in enumerate(block.split('\n')):
                line = line.strip()
                if not line: continue
                if i > 0: p.add_run().add_break()
                parse_inline_markdown(p, line, images=_images_dict,
                                      content_width_cm=content_width_cm)  # ПРАВКА #32
            last_regular_paragraph = p
            # ПРАВКА #20: лид-абзац (целиком жирный) держится со следующим блоком
            stripped_block = block.strip()
            if (stripped_block.startswith('**')
                    and stripped_block.endswith('**')
                    and stripped_block.count('**') == 2):
                set_keep_with_next(p)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(2)
            after_heading = False

    doc.save(output_filename)
    print(f"✅ Готово! Файл сохранён: {output_filename}")


# =============================================================================
# ЗАПУСК
# =============================================================================
if __name__ == '__main__':
    if not os.path.exists(INPUT_FILE):
        print(f"❌ .md файл не найден: {INPUT_FILE}")
        print("   Проверь путь INPUT_FILE в начале скрипта.")
        exit(1)

    if not os.path.exists(TEMPLATE_FILE):
        print(f"⚠️  Шаблон не найден: {TEMPLATE_FILE}")
        print("   Документ будет создан без фирменного хедера.\n")

    print(f"Конвертирую: {os.path.basename(INPUT_FILE)}")
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        md_text = f.read()

    template = TEMPLATE_FILE if os.path.exists(TEMPLATE_FILE) else None
    convert_md_to_docx(md_text, OUTPUT_FILE, template_path=template)
